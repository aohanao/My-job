import streamlit as st
import os
import subprocess
from knowledge_base import get_file_md5, check_md5, save_md5

# --- 页面配置 ---
st.set_page_config(page_title="CAE 知识库批量更新", page_icon="📚", layout="wide")
st.title("📚 CAE 专属知识库批量更新控制台")
st.markdown("支持**多文件同时上传** (PDF / Markdown)，系统将自动排队、深度解析并注入大模型知识库。")

# --- 1. 定义工作目录 ---
BASE_DIR = "G:/vscode/LangChain_Project/CAE_RAG_project" # 部署到服务器时，请确保修改为服务器的绝对路径！
TEMP_UPLOAD_DIR = os.path.join(BASE_DIR, "temp_uploads") 
MARKER_OUTPUT_DIR = os.path.join(BASE_DIR, "marker_output") 
MAX_FILES_LIMIT = 5  # 单次上传最大文件限制

for d in [TEMP_UPLOAD_DIR, MARKER_OUTPUT_DIR]:
    os.makedirs(d, exist_ok=True)

# ==========================================
# 🌟 全局缓存锁 (懒加载)
# ==========================================
@st.cache_resource
def get_kb_service():
    """
    使用 Streamlit 的资源缓存，确保全局只有一个 KnowledgeBaseService 实例。
    防止重复加载沉重的 Embedding 模型，提升系统响应速度。
    """
    from knowledge_base import KnowledgeBaseService
    return KnowledgeBaseService()

# --- 2. 界面布局 (使用选项卡) ---
tab_upload, tab_manage = st.tabs(["📤 批量文件上传", "📂 知识库现有管理"])

with tab_upload:
    st.subheader("文件上传队列")
    uploaded_files = st.file_uploader(
        f"请将需要入库的 CAE 文献拖拽至此处 (支持 pdf/md/docx/txt，单次最多允许 {MAX_FILES_LIMIT} 个)", 
        type=["pdf", "md", "docx", "txt"],
        accept_multiple_files=True 
    )

    if uploaded_files:
        total_files = len(uploaded_files)
        
        if total_files > MAX_FILES_LIMIT:
            st.error(f"❌ 队列超载！您当前选中了 **{total_files}** 个文件，系统单次最多支持处理 **{MAX_FILES_LIMIT}** 个。")
            st.warning("👉 请点击上方文件列表右侧的 'X' 移除多余的文档，以解锁入库功能。")
        else:
            total_size_mb = sum([f.size for f in uploaded_files]) / (1024 * 1024)
            st.info(f"📁 待处理队列合规：共 **{total_files}** 个文件，总大小 **{total_size_mb:.2f} MB**")

            if st.button("🚀 开始批量解析与入库", type="primary"):        
                with st.spinner("⏳ 正在唤醒大模型与向量数据库引擎..."):
                    kb_service = get_kb_service()

                progress_bar = st.progress(0.0, text="准备启动批处理队列...")
                st.markdown("### 📝 实时执行日志")
                log_container = st.container(height=450)
            
                stats = {"success": 0, "skip": 0, "error": 0}
                error_details = []

                for index, file in enumerate(uploaded_files):
                    file_name = file.name
                    file_type = file_name.split(".")[-1].lower()
                    base_name = os.path.splitext(file_name)[0]
                    file_bytes = file.getvalue()

                    file_md5 = get_file_md5(file_bytes)
                    if check_md5(file_md5):
                        log_container.warning(f"⏭️ [秒传拦截] '{file_name}' 物理指纹已存在，瞬间跳过。")
                        stats["skip"] += 1
                        continue
                    
                    current_progress = (index) / total_files
                    progress_bar.progress(current_progress, text=f"正在处理 ({index+1}/{total_files}): {file_name}")
                
                    if index > 0:
                        log_container.markdown(f"---")
                    log_container.write(f"▶️ **开始处理文档 [{index+1}/{total_files}]:** `{file_name}`")
                
                    temp_file_path = os.path.join(TEMP_UPLOAD_DIR, file_name)
                    marker_folder_path = os.path.join(MARKER_OUTPUT_DIR, base_name)
                    marker_md_path = os.path.join(marker_folder_path, f"{base_name}.md")
                    parsed_text = None
                
                    try:
                        with open(temp_file_path, "wb") as f:
                            f.write(file.getvalue())

                        if file_type == "md":
                            with open(temp_file_path, "r", encoding="utf-8") as f:
                                parsed_text = f.read()
                            log_container.write("✅ Markdown 读取成功")
                        
                        elif file_type == "docx":
                            log_container.warning("📑 正在启动深度解析引擎 (Unstructured)...")
                            try:
                                # 尝试 1: 使用 unstructured 进行解析（支持更复杂的结构）
                                from unstructured.partition.docx import partition_docx
                                elements = partition_docx(filename=temp_file_path)
                                parsed_text = "\n\n".join([str(el) for el in elements])
                                
                                # 尝试 2: 如果结果为空，使用 python-docx 强行提取文字（保底方案）
                                if not parsed_text.strip():
                                    log_container.info("⚠️ 深度解析返回为空，正在切换至 python-docx 兼容模式...")
                                    import docx
                                    doc = docx.Document(temp_file_path)
                                    # 1. 抓取正文段落
                                    all_texts = [para.text for para in doc.paragraphs if para.text.strip()]
                                    # 2. 深度抓取所有表格中的文字
                                    for table in doc.tables:
                                        for row in table.rows:
                                            for cell in row.cells:
                                                if cell.text.strip():
                                                    all_texts.append(cell.text.strip())
                                    seen = set()
                                    unique_texts = [x for x in all_texts if not (x in seen or seen.add(x))]
                                    parsed_text = "\n".join(unique_texts)
                                
                                if not parsed_text.strip():
                                    raise Exception("Word 文档解析结果为空，请确认文档中是否有可编辑文本。")
                                    
                                log_container.write(f"✅ Word 解析成功 (字数: {len(parsed_text)})")
                            except ImportError:
                                raise Exception("未检测到必要的 Word 处理库，请执行 `pip install unstructured python-docx`。")
                            except Exception as e:
                                raise Exception(f"Word 解析引擎故障: {str(e)}")

                        elif file_type == "txt":
                            log_container.write("📄 正在读取文本文件...")
                            parsed_text = None
                            for encoding in ["utf-8", "gbk", "gb2312", "utf-16"]:
                                try:
                                    with open(temp_file_path, "r", encoding=encoding) as f:
                                        parsed_text = f.read()
                                    log_container.write(f"✅ TXT 读取成功 (编码: {encoding})")
                                    break
                                except Exception:
                                    continue
                            if parsed_text is None:
                                raise Exception("无法识别的文本编码格式")

                        elif file_type == "pdf":
                            log_container.write("🧠 正在调用 Marker 提取 PDF 内容与公式...")
                            command = ["marker_single", temp_file_path, "--output_dir", MARKER_OUTPUT_DIR]
                            process = subprocess.Popen(
                                command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, 
                                text=True, bufsize=1, encoding="utf-8", errors="replace"
                            )
                            for line in process.stdout:
                                print(line, end="", flush=True) 
                            process.wait()
                        
                            if process.returncode == 0 and os.path.exists(marker_md_path):
                                with open(marker_md_path, "r", encoding="utf-8") as f:
                                    parsed_text = f.read()
                                log_container.write("✅ PDF 深度解析成功")
                            else:
                                raise Exception("Marker 解析失败")
                        
                        else:
                            log_container.error(f"❌ '{file_name}' 是不支持的文件格式，目前仅支持 pdf/md/docx/txt。")
                            raise Exception(f"不支持的后缀: {file_type}")

                        if parsed_text:
                            log_container.write("⚙️ 正在切片与多模态分析...")
                            result = kb_service.upload_by_str(
                                data=parsed_text, 
                                filename=file_name, 
                                file_md5=file_md5,  # 👈 传入 MD5
                                base_dir=marker_folder_path if file_type == "pdf" else os.path.dirname(temp_file_path),
                                progress_callback=log_container.write
                            )
                        
                            if "[成功]" in result:
                                stats["success"] += 1
                                log_container.success(f"🎉 入库完成：{result}")
                                save_md5(file_md5)
                            else:
                                raise Exception(result)

                    except Exception as e:
                        stats["error"] += 1
                        error_msg = f"❌ '{file_name}' 处理失败: {str(e)}"
                        log_container.error(error_msg)
                        error_details.append(error_msg)
                    
                    finally:
                        if os.path.exists(temp_file_path):
                            os.remove(temp_file_path)

                progress_bar.progress(1.0, text="✅ 队列处理完毕！")
                st.balloons()
            
                st.markdown("### 📊 批量入库战报")
                col1, col2, col3, col4 = st.columns(4)
                col1.metric("总计提交", f"{total_files} 份")
                col2.metric("✅ 成功入库", f"{stats['success']} 份")
                col3.metric("⏭️ 跳过重复", f"{stats['skip']} 份")
                col4.metric("❌ 解析失败", f"{stats['error']} 份")
            
                if error_details:
                    with st.expander("查看失败详情"):
                        for err in error_details:
                            st.error(err)

with tab_manage:
    st.subheader("📂 现有文档管理")
    kb_service = get_kb_service()
    
    # 刷新按钮
    if st.button("🔄 刷新文件列表"):
        st.rerun()

    files = kb_service.get_all_sources()
    
    if not files:
        st.info("目前知识库中还没有任何文档。")
    else:
        st.write(f"当前共入库 **{len(files)}** 个独立文档：")
        
        # 使用表格展示，并添加删除逻辑
        for file_info in files:
            with st.container(border=True):
                col1, col2, col3 = st.columns([4, 2, 1])
                with col1:
                    st.markdown(f"**📄 {file_info['filename']}**")
                    st.caption(f"入库时间: {file_info['create_time']} | 切片数量: {file_info['chunk_count']}")
                with col2:
                    if file_info['file_md5']:
                        st.code(file_info['file_md5'][:16], language=None)
                    else:
                        st.caption("📜 [历史遗留文档]")
                with col3:
                    if st.button("🗑️ 删除", key=f"del_{file_info['filename']}"):
                        with st.spinner(f"正在移除 {file_info['filename']}..."):
                            if kb_service.delete_by_source(file_info['filename']):
                                st.success(f"已成功删除 {file_info['filename']}")
                                time_import = __import__("time")
                                time_import.sleep(1)
                                st.rerun()
                            else:
                                st.error("删除失败，请查看后台日志。")